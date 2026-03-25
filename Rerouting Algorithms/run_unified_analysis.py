"""
Unified Analysis Script
Runs baseline, rerouting, and naive rerouting for a set of seeds (default 20),
extracts seven metrics from each run, and writes a combined CSV, optional Excel, optional Word doc,
and optional markdown file with full terminal output.
"""

import sys
import os
import csv
import json
import argparse
import glob
import re
from typing import List, Optional, Dict, Any, Tuple

# Default 20 seeds for unified analysis
DEFAULT_SEEDS = list(range(1001, 1021))  # 1001 .. 1020


def extract_seven_metrics(stats: Optional[Dict[str, Any]], from_baseline_save: bool = False) -> Dict[str, Any]:
    """
    Extract the seven unified metrics from a stats dict (flat or nested under 'statistics').
    Returns a normalized dict: average_delay, crossing_delay, queue_length, los,
    speed_ratio_efficiency, vehicles_affected, total_delay.
    """
    if stats is None:
        return {
            "average_delay": None,
            "crossing_delay": None,
            "queue_length": None,
            "los": None,
            "speed_ratio_efficiency": None,
            "vehicles_affected": None,
            "total_delay": None,
        }
    if from_baseline_save and "statistics" in stats:
        stats = stats["statistics"]
    # Derive queue_length_max from crossing_stats if missing (older JSONs)
    queue_length = stats.get("queue_length_max")
    if queue_length is None and stats.get("crossing_stats"):
        lengths = [c.get("max_queue_length", 0) for c in stats["crossing_stats"].values()]
        queue_length = max(lengths, default=0)
    # Derive total_delay_all if missing (older JSONs)
    total_delay = stats.get("total_delay_all")
    if total_delay is None:
        n = stats.get("total_vehicles_seen") or stats.get("total_vehicles")
        avg = stats.get("avg_delay_time_all")
        if n is not None and avg is not None:
            total_delay = avg * n
    return {
        "average_delay": stats.get("avg_delay_time_all"),
        "crossing_delay": stats.get("avg_crossing_delay_all"),
        "queue_length": queue_length,
        "los": stats.get("los_all"),
        "speed_ratio_efficiency": stats.get("speed_ratio_all"),
        "vehicles_affected": stats.get("vehicles_affected_by_crossings"),
        "total_delay": total_delay,
    }


def _cell_value(val: Any) -> str:
    """Format a value for table output (Word/Excel)."""
    if val is None:
        return ""
    if isinstance(val, float):
        return f"{val:.4g}" if val == val else ""  # avoid nan
    return str(val)


def run_unified_analysis(
    seeds: Optional[List[int]] = None,
    sumo_binary: str = "sumo",
    simulation_end_time: float = 3600,
    output_csv: str = "unified_analysis_20seeds.csv",
    output_excel: Optional[str] = "unified_analysis_20seeds.xlsx",
    output_doc: Optional[str] = "unified_analysis_20seeds.docx",
) -> List[Dict[str, Any]]:
    """
    Run baseline, rerouting, and naive rerouting for each seed; collect seven metrics;
    write CSV, optional Excel, and optional Word doc.
    Returns the list of rows (each with seed, scenario, and the seven metrics).
    """
    from run_baseline_analysis import run_baseline_analysis
    from run_rerouting_analysis import run_rerouting_analysis
    from run_naive_rerouting_analysis import run_naive_rerouting_analysis

    if seeds is None:
        seeds = DEFAULT_SEEDS

    rows: List[Dict[str, Any]] = []
    for seed in seeds:
        for scenario_name, run_fn in [
            ("baseline", run_baseline_analysis),
            ("rerouting", run_rerouting_analysis),
            ("naive", run_naive_rerouting_analysis),
        ]:
            print("\n" + "=" * 70)
            print(f"Unified run: seed={seed} scenario={scenario_name}")
            print("=" * 70)
            try:
                stats = run_fn(
                    seed=seed,
                    sumo_binary=sumo_binary,
                    simulation_end_time=simulation_end_time,
                    generate_visuals=False,
                )
                seven = extract_seven_metrics(stats, from_baseline_save=False)
                row = {"seed": seed, "scenario": scenario_name, **seven}
                rows.append(row)
            except Exception as e:
                print(f"[WARNING] {scenario_name} seed={seed} failed: {e}")
                row = {
                    "seed": seed,
                    "scenario": scenario_name,
                    "average_delay": None,
                    "crossing_delay": None,
                    "queue_length": None,
                    "los": None,
                    "speed_ratio_efficiency": None,
                    "vehicles_affected": None,
                    "total_delay": None,
                }
                rows.append(row)

    # Write CSV
    fieldnames = [
        "seed", "scenario",
        "average_delay", "crossing_delay", "queue_length", "los",
        "speed_ratio_efficiency", "vehicles_affected", "total_delay",
    ]
    with open(output_csv, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        w.writeheader()
        w.writerows(rows)
    print(f"\n[OK] CSV written to {output_csv}")

    # Optional Excel (openpyxl required for Excel output)
    if output_excel:
        try:
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Unified Analysis"
            ws.append(fieldnames)
            for r in rows:
                ws.append([r.get(f) for f in fieldnames])
            wb.save(output_excel)
            print(f"[OK] Excel written to {output_excel}")
        except ImportError:
            print(f"[WARNING] Excel export skipped: install openpyxl for Excel output")
        except Exception as ex:
            print(f"[WARNING] Excel export skipped: {ex}")

    # Optional Word document (python-docx required for .docx output)
    if output_doc:
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Unified Analysis Results", level=0)
            doc.add_paragraph(
                f"Seeds: {min(seeds)}–{max(seeds)} ({len(seeds)} seeds). "
                f"Scenarios: baseline, rerouting, naive. Metrics: average_delay, crossing_delay, "
                "queue_length, LOS, speed_ratio_efficiency, vehicles_affected, total_delay."
            )
            table = doc.add_table(rows=1 + len(rows), cols=len(fieldnames))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, name in enumerate(fieldnames):
                hdr_cells[i].text = name.replace("_", " ").title()
            for r_idx, r in enumerate(rows):
                row_cells = table.rows[r_idx + 1].cells
                for c_idx, f in enumerate(fieldnames):
                    row_cells[c_idx].text = _cell_value(r.get(f))
            doc.save(output_doc)
            print(f"[OK] Word document written to {output_doc}")
        except ImportError:
            print(f"[WARNING] Word export skipped: install python-docx for .docx output")
        except Exception as ex:
            print(f"[WARNING] Word export skipped: {ex}")

    return rows


def build_unified_results_from_jsons(
    json_dir: str,
    output_csv: str = "unified_analysis_from_jsons.csv",
    output_excel: Optional[str] = "unified_analysis_from_jsons.xlsx",
    output_doc: Optional[str] = "unified_analysis_from_jsons.docx",
) -> List[Dict[str, Any]]:
    """
    Build unified results table from existing stats_*.json files in json_dir.
    Does not run simulations. Writes CSV (and optionally Excel/Word).
    """
    # stats_baseline_seed1001.json -> (1001, baseline), etc.
    pattern = re.compile(r"stats_(baseline|rerouting|naive_rerouting)(?:_seed(\d+))?\.json$", re.I)
    rows: List[Dict[str, Any]] = []
    seen: set = set()

    for path in glob.glob(os.path.join(json_dir, "stats_*.json")):
        name = os.path.basename(path)
        m = pattern.match(name)
        if not m:
            continue
        scenario = m.group(1).replace("naive_rerouting", "naive")
        seed_str = m.group(2)
        seed = int(seed_str) if seed_str else None
        if (seed, scenario) in seen:
            continue
        seen.add((seed, scenario))
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            print(f"[WARNING] Skip {name}: {e}")
            continue
        from_baseline = scenario == "baseline" and "statistics" in data
        seven = extract_seven_metrics(data, from_baseline_save=from_baseline)
        row = {"seed": seed, "scenario": scenario, **seven}
        rows.append(row)

    rows.sort(key=lambda r: (r.get("seed") or 0, r.get("scenario", "")))
    fieldnames = [
        "seed", "scenario",
        "average_delay", "crossing_delay", "queue_length", "los",
        "speed_ratio_efficiency", "vehicles_affected", "total_delay",
    ]
    with open(output_csv, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        w.writeheader()
        w.writerows(rows)
    print(f"[OK] CSV written to {output_csv} ({len(rows)} rows)")

    if output_excel:
        try:
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Unified Analysis"
            ws.append(fieldnames)
            for r in rows:
                ws.append([r.get(f) for f in fieldnames])
            wb.save(output_excel)
            print(f"[OK] Excel written to {output_excel}")
        except Exception as ex:
            print(f"[WARNING] Excel export skipped: {ex}")
    if output_doc:
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Unified Analysis Results (from JSON)", level=0)
            doc.add_paragraph(f"Source directory: {os.path.abspath(json_dir)}. Rows: {len(rows)}.")
            table = doc.add_table(rows=1 + len(rows), cols=len(fieldnames))
            table.style = "Table Grid"
            for i, name in enumerate(fieldnames):
                table.rows[0].cells[i].text = name.replace("_", " ").title()
            for r_idx, r in enumerate(rows):
                for c_idx, f in enumerate(fieldnames):
                    table.rows[r_idx + 1].cells[c_idx].text = _cell_value(r.get(f))
            doc.save(output_doc)
            print(f"[OK] Word document written to {output_doc}")
        except Exception as ex:
            print(f"[WARNING] Word export skipped: {ex}")
    return rows


def _parse_seeds(seeds_arg: Optional[str]) -> List[int]:
    """Parse --seeds into a list of ints. If missing, return DEFAULT_SEEDS."""
    if not seeds_arg:
        return DEFAULT_SEEDS
    cleaned = seeds_arg.replace(",", " ").split()
    return [int(x) for x in cleaned if x.strip()]


class _TeeLogger:
    """Write to both the real terminal and an internal buffer (for saving terminal output to .md)."""
    def __init__(self, stream):
        self._stream = stream
        self._lines: List[str] = []

    def write(self, msg: str) -> None:
        self._stream.write(msg)
        self._lines.append(msg)

    def flush(self) -> None:
        self._stream.flush()

    def get_text(self) -> str:
        return "".join(self._lines)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Run baseline, rerouting, and naive rerouting for multiple seeds and export unified metrics."
    )
    parser.add_argument(
        "--seeds",
        type=str,
        default=None,
        help=f"Comma/space separated seeds (default: 1001..1020, i.e. {DEFAULT_SEEDS[0]} to {DEFAULT_SEEDS[-1]}).",
    )
    parser.add_argument(
        "--sumo-binary",
        type=str,
        choices=["sumo", "sumo-gui"],
        default="sumo",
        help="SUMO binary (default: sumo for batch).",
    )
    parser.add_argument("--end-time", type=float, default=3600, help="Simulation end time in seconds.")
    parser.add_argument("--output-csv", type=str, default="unified_analysis_20seeds.csv", help="Output CSV path.")
    parser.add_argument("--output-excel", type=str, default="unified_analysis_20seeds.xlsx", help="Output Excel path (set to empty to skip).")
    parser.add_argument("--output-doc", type=str, default="unified_analysis_20seeds.docx", help="Output Word document path (set to empty to skip).")
    parser.add_argument("--output-md", type=str, default=None, help="Output markdown path for full terminal output (default: Unified_Analysis_Terminal_Output_Seeds_<seeds>.md). Set to empty to skip.")
    parser.add_argument("--from-json-dir", type=str, default=None, help="Build unified table from existing stats_*.json in this directory (no simulation). Example: ../Testing")
    args = parser.parse_args()

    if args.from_json_dir:
        build_unified_results_from_jsons(
            args.from_json_dir,
            output_csv=args.output_csv,
            output_excel=args.output_excel or None,
            output_doc=args.output_doc or None,
        )
        sys.exit(0)

    seeds = _parse_seeds(args.seeds)
    output_excel = args.output_excel if args.output_excel else None
    output_doc = args.output_doc if args.output_doc else None
    output_md = args.output_md if args.output_md else None
    if output_md is None and args.output_md != "":
        # Default: build filename from seeds like Rerouting script
        seed_labels = [str(s) for s in seeds]
        output_md = f"Unified_Analysis_Terminal_Output_Seeds_{'_'.join(seed_labels)}.md"

    # Tee stdout so we can save full terminal output to .md
    _original_stdout = sys.stdout
    if output_md:
        _tee = _TeeLogger(_original_stdout)
        sys.stdout = _tee

    run_unified_analysis(
        seeds=seeds,
        sumo_binary=args.sumo_binary,
        simulation_end_time=args.end_time,
        output_csv=args.output_csv,
        output_excel=output_excel,
        output_doc=output_doc,
    )

    # Restore stdout and save terminal output to markdown
    if output_md:
        sys.stdout = _original_stdout
        captured = _tee.get_text()
        with open(output_md, "w", encoding="utf-8") as f:
            f.write("# Unified analysis – full terminal output\n\n")
            f.write(f"**Seeds:** {', '.join(str(s) for s in seeds)}\n\n")
            f.write("**Scenarios:** baseline, rerouting, naive\n\n")
            f.write("---\n\n```\n")
            f.write(captured)
            f.write("\n```\n")
        print(f"\n[OK] Terminal output saved to {output_md}")
